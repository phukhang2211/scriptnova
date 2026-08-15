"""Export helpers — generate DOCX and SRT files from a completed Job."""

from __future__ import annotations

import re
from io import BytesIO
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from .models import Job


# ── SRT ──────────────────────────────────────────────────────────────────────

def _ms_to_srt_time(ms: int | float) -> str:
    """Convert milliseconds to SRT timestamp format HH:MM:SS,mmm."""
    ms = int(ms or 0)
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1_000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def generate_srt(job: Job) -> str:
    """Return SRT-formatted string. Uses utterances when available, falls back to transcript."""
    if job.utterances:
        blocks = []
        for i, sentence in enumerate(job.utterances, 1):
            start = _ms_to_srt_time(sentence.get("start", 0))
            end   = _ms_to_srt_time(sentence.get("end", 0))
            blocks.append(f"{i}\n{start} --> {end}\n{sentence['text']}\n")
        return "\n".join(blocks)

    # Fallback: split transcript into lines, assign sequential fake timestamps
    transcript = (job.transcript or "").strip()
    if not transcript:
        return ""
    lines = [l.strip() for l in transcript.splitlines() if l.strip()]
    blocks = []
    for i, line in enumerate(lines, 1):
        start = _ms_to_srt_time((i - 1) * 4000)
        end   = _ms_to_srt_time(i * 4000)
        blocks.append(f"{i}\n{start} --> {end}\n{line}\n")
    return "\n".join(blocks)


# ── DOCX ─────────────────────────────────────────────────────────────────────

def generate_docx(job: Job) -> bytes:
    """Return DOCX file as bytes. Falls back gracefully when utterances are missing."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    base_name = job.original_filename.rsplit(".", 1)[0] if job.original_filename else f"Job #{job.pk}"
    heading = doc.add_heading(base_name, level=1)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # Metadata line
    meta_parts = []
    if job.audio_duration_seconds:
        m, s = divmod(int(job.audio_duration_seconds), 60)
        meta_parts.append(f"Duration: {m}m {s:02d}s")
    if job.created_at:
        meta_parts.append(job.created_at.strftime("Transcribed: %d %b %Y"))
    if meta_parts:
        meta = doc.add_paragraph(" · ".join(meta_parts))
        for run in meta.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    doc.add_paragraph()  # spacer

    # Body — one paragraph per sentence if we have utterances, else best-effort split
    if job.utterances:
        for sentence in job.utterances:
            doc.add_paragraph(sentence["text"])
    else:
        # Split on sentence endings as a fallback
        for chunk in re.split(r"(?<=[.!?])\s+", job.transcript):
            chunk = chunk.strip()
            if chunk:
                doc.add_paragraph(chunk)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
